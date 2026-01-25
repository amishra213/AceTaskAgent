"""
Data Extraction Sub-Agent for analyzing and extracting relevant information
from input and temp folder data.

This agent intelligently extracts only the relevant portions of input data
based on the current task objective, avoiding large context overhead.

Capabilities:
- Analyze input files and extract task-relevant information
- Summarize large documents to key points
- Filter and extract specific data based on queries
- Cache extracted data for reuse within a session
- Prioritize and rank input files by relevance to task
"""

from typing import Optional, List, Dict, Any, Union, Tuple
from pathlib import Path
from datetime import datetime
import json
import re
import hashlib
from dataclasses import dataclass, field

from task_manager.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class ExtractionResult:
    """Result of data extraction from a single file."""
    file_path: str
    file_name: str
    file_type: str
    relevance_score: float  # 0.0 to 1.0
    extracted_content: str
    key_points: List[str]
    metadata: Dict[str, Any]
    extraction_time: datetime = field(default_factory=datetime.now)
    truncated: bool = False
    original_size: int = 0
    extracted_size: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "file_path": self.file_path,
            "file_name": self.file_name,
            "file_type": self.file_type,
            "relevance_score": self.relevance_score,
            "extracted_content": self.extracted_content,
            "key_points": self.key_points,
            "metadata": self.metadata,
            "extraction_time": self.extraction_time.isoformat(),
            "truncated": self.truncated,
            "original_size": self.original_size,
            "extracted_size": self.extracted_size
        }


@dataclass
class ExtractionContext:
    """Context for a data extraction session."""
    objective: str
    keywords: List[str]
    max_content_per_file: int = 5000  # Characters
    max_total_content: int = 50000   # Characters
    include_file_types: Optional[List[str]] = None
    exclude_patterns: Optional[List[str]] = None
    relevance_threshold: float = 0.1  # Minimum relevance to include


class DataExtractionAgent:
    """
    Sub-agent for intelligently extracting relevant information from input/temp files.
    
    This agent:
    1. Analyzes files in input and temp folders
    2. Scores file relevance based on the current task objective
    3. Extracts only relevant portions of content
    4. Summarizes large documents to key points
    5. Caches extractions to avoid redundant processing
    """
    
    # Supported file types for extraction
    SUPPORTED_TYPES = {
        'text': ['.txt', '.md', '.rst', '.log'],
        'data': ['.json', '.csv', '.xml', '.yaml', '.yml'],
        'pdf': ['.pdf'],
        'excel': ['.xlsx', '.xls', '.xlsm', '.csv'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'],
        'document': ['.doc', '.docx', '.rtf', '.odt']
    }
    
    # Maximum sizes for different operations
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    MAX_TEXT_PREVIEW = 10000  # Characters for text preview
    MAX_SUMMARY_LINES = 50
    
    def __init__(
        self,
        llm_client: Optional[Any] = None,
        cache_extractions: bool = True
    ):
        """
        Initialize Data Extraction Agent.
        
        Args:
            llm_client: LLM client for intelligent summarization
            cache_extractions: Whether to cache extraction results
        """
        self.llm_client = llm_client
        self.cache_extractions = cache_extractions
        self._extraction_cache: Dict[str, ExtractionResult] = {}
        self._session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        self.supported_operations = [
            "extract_relevant",
            "summarize_file",
            "search_content",
            "rank_by_relevance",
            "extract_structured",
            "get_file_preview"
        ]
        
        # Check for optional dependencies
        self._check_dependencies()
        logger.info("Data Extraction Agent initialized")
    
    def _check_dependencies(self):
        """Check for optional dependencies."""
        self.has_pdf = False
        self.has_excel = False
        self.has_docx = False
        
        try:
            import PyPDF2  # noqa: F401
            self.has_pdf = True
        except ImportError:
            try:
                import pypdf  # noqa: F401
                self.has_pdf = True
            except ImportError:
                logger.debug("PDF support not available")
        
        try:
            import openpyxl  # noqa: F401
            self.has_excel = True
        except ImportError:
            logger.debug("Excel support not available")
        
        try:
            from docx import Document  # noqa: F401
            self.has_docx = True
        except ImportError:
            logger.debug("DOCX support not available")
    
    def extract_relevant_data(
        self,
        input_folder: Union[str, Path],
        objective: str,
        keywords: Optional[List[str]] = None,
        temp_folder: Optional[Union[str, Path]] = None,
        max_files: int = 20,
        max_content_per_file: int = 5000,
        max_total_content: int = 50000,
        file_types: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Extract relevant data from input (and optionally temp) folder.
        
        This is the main entry point for the agent. It:
        1. Scans folders for files
        2. Ranks files by relevance to objective
        3. Extracts relevant content from top files
        4. Returns a context-optimized data package
        
        Args:
            input_folder: Path to input folder
            objective: The task objective for relevance scoring
            keywords: Additional keywords for relevance matching
            temp_folder: Optional temp folder to also scan
            max_files: Maximum number of files to process
            max_content_per_file: Max characters per file
            max_total_content: Max total characters in result
            file_types: Filter to specific file types
            
        Returns:
            Dictionary with extracted data and metadata
        """
        logger.info(f"[DATA_EXTRACT] Starting extraction for objective: {objective[:100]}...")
        
        input_path = Path(input_folder).resolve()
        temp_path = Path(temp_folder).resolve() if temp_folder else None
        
        # Build extraction context
        if keywords is None:
            keywords = self._extract_keywords(objective)
        
        context = ExtractionContext(
            objective=objective,
            keywords=keywords,
            max_content_per_file=max_content_per_file,
            max_total_content=max_total_content,
            include_file_types=file_types
        )
        
        # Scan and collect files
        all_files = self._scan_folders(input_path, temp_path, file_types)
        logger.info(f"[DATA_EXTRACT] Found {len(all_files)} files to analyze")
        
        if not all_files:
            return {
                "success": True,
                "files_found": 0,
                "files_processed": 0,
                "extractions": [],
                "summary": "No files found in input folder",
                "total_content_size": 0
            }
        
        # Rank files by relevance
        ranked_files = self._rank_files_by_relevance(all_files, context)
        
        # Extract from top files up to limits
        extractions: List[ExtractionResult] = []
        total_content_size = 0
        files_processed = 0
        
        for file_path, relevance_score in ranked_files[:max_files]:
            if total_content_size >= max_total_content:
                logger.debug(f"Reached max total content limit ({max_total_content} chars)")
                break
            
            if relevance_score < context.relevance_threshold:
                logger.debug(f"Skipping {file_path.name} (relevance {relevance_score:.2f} below threshold)")
                continue
            
            # Check cache
            cache_key = self._get_cache_key(file_path, context)
            if self.cache_extractions and cache_key in self._extraction_cache:
                extraction = self._extraction_cache[cache_key]
                logger.debug(f"Using cached extraction for {file_path.name}")
            else:
                # Extract content
                remaining_budget = max_total_content - total_content_size
                extraction = self._extract_from_file(
                    file_path,
                    context,
                    max_chars=min(max_content_per_file, remaining_budget)
                )
                extraction.relevance_score = relevance_score
                
                if self.cache_extractions:
                    self._extraction_cache[cache_key] = extraction
            
            extractions.append(extraction)
            total_content_size += extraction.extracted_size
            files_processed += 1
        
        # Build summary
        summary = self._build_extraction_summary(extractions, context)
        
        result = {
            "success": True,
            "objective": objective,
            "keywords": keywords,
            "files_found": len(all_files),
            "files_processed": files_processed,
            "extractions": [e.to_dict() for e in extractions],
            "summary": summary,
            "total_content_size": total_content_size,
            "extraction_time": datetime.now().isoformat()
        }
        
        logger.info(f"[DATA_EXTRACT] Processed {files_processed} files, extracted {total_content_size} chars")
        return result
    
    def _extract_keywords(self, objective: str) -> List[str]:
        """Extract keywords from objective text."""
        # Remove common stop words and extract meaningful terms
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
            'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
            'could', 'should', 'may', 'might', 'must', 'shall', 'can', 'need',
            'that', 'this', 'these', 'those', 'it', 'its', 'what', 'which', 'who',
            'whom', 'how', 'when', 'where', 'why', 'all', 'each', 'every', 'both',
            'few', 'more', 'most', 'other', 'some', 'such', 'than', 'too', 'very',
            'just', 'about', 'into', 'through', 'during', 'before', 'after',
            'above', 'below', 'between', 'under', 'again', 'further', 'then',
            'once', 'here', 'there', 'any', 'only', 'same', 'so', 'as', 'if'
        }
        
        # Tokenize and filter
        words = re.findall(r'\b[a-zA-Z]{3,}\b', objective.lower())
        keywords = [w for w in words if w not in stop_words]
        
        # Get unique keywords, preserve order, limit count
        seen = set()
        unique_keywords = []
        for kw in keywords:
            if kw not in seen:
                seen.add(kw)
                unique_keywords.append(kw)
        
        return unique_keywords[:20]
    
    def _scan_folders(
        self,
        input_folder: Path,
        temp_folder: Optional[Path],
        file_types: Optional[List[str]]
    ) -> List[Path]:
        """Scan folders and return list of file paths."""
        files = []
        
        folders_to_scan = [input_folder]
        if temp_folder and temp_folder.exists():
            folders_to_scan.append(temp_folder)
        
        # Build extension filter
        allowed_extensions = set()
        if file_types:
            for ft in file_types:
                if ft in self.SUPPORTED_TYPES:
                    allowed_extensions.update(self.SUPPORTED_TYPES[ft])
        else:
            for exts in self.SUPPORTED_TYPES.values():
                allowed_extensions.update(exts)
        
        for folder in folders_to_scan:
            if not folder.exists():
                continue
                
            for root, dirs, filenames in os.walk(folder):
                # Skip hidden directories
                dirs[:] = [d for d in dirs if not d.startswith('.')]
                
                for filename in filenames:
                    if filename.startswith('.'):
                        continue
                    
                    file_path = Path(root) / filename
                    ext = file_path.suffix.lower()
                    
                    if ext in allowed_extensions:
                        # Check file size
                        try:
                            if file_path.stat().st_size <= self.MAX_FILE_SIZE:
                                files.append(file_path)
                        except OSError:
                            continue
        
        return files
    
    def _rank_files_by_relevance(
        self,
        files: List[Path],
        context: ExtractionContext
    ) -> List[Tuple[Path, float]]:
        """
        Rank files by relevance to the objective.
        
        Uses multiple signals:
        1. Filename matching keywords
        2. File extension/type relevance
        3. File size (prefer medium-sized files)
        4. Content preview matching (for text files)
        """
        scored_files: List[Tuple[Path, float]] = []
        keywords_lower = [kw.lower() for kw in context.keywords]
        
        for file_path in files:
            score = 0.0
            
            # Filename relevance (40% weight)
            filename_lower = file_path.stem.lower()
            filename_words = set(re.findall(r'\b[a-zA-Z]+\b', filename_lower))
            
            keyword_matches = sum(1 for kw in keywords_lower if kw in filename_lower)
            word_matches = len(filename_words & set(keywords_lower))
            
            if keyword_matches > 0:
                score += 0.4 * min(keyword_matches / len(keywords_lower), 1.0)
            if word_matches > 0:
                score += 0.2 * min(word_matches / len(keywords_lower), 1.0)
            
            # File type relevance (20% weight)
            ext = file_path.suffix.lower()
            if ext in ['.pdf', '.docx', '.doc']:
                score += 0.15  # Document files likely contain useful info
            elif ext in ['.csv', '.xlsx', '.xls']:
                score += 0.20  # Data files often very relevant
            elif ext in ['.json', '.xml', '.yaml']:
                score += 0.15  # Structured data
            elif ext in ['.txt', '.md']:
                score += 0.10  # Text files
            
            # File size scoring (10% weight) - prefer medium files
            try:
                size = file_path.stat().st_size
                if 1000 <= size <= 500000:  # 1KB to 500KB is ideal
                    score += 0.10
                elif 500 <= size <= 1000000:  # 500B to 1MB acceptable
                    score += 0.05
            except OSError:
                pass
            
            # Content preview for text files (30% weight)
            if ext in ['.txt', '.md', '.csv', '.json', '.xml', '.yaml', '.yml', '.log']:
                try:
                    preview = self._read_text_preview(file_path, 2000)
                    preview_lower = preview.lower()
                    content_matches = sum(1 for kw in keywords_lower if kw in preview_lower)
                    if content_matches > 0:
                        score += 0.3 * min(content_matches / len(keywords_lower), 1.0)
                except Exception:
                    pass
            
            scored_files.append((file_path, score))
        
        # Sort by score descending
        scored_files.sort(key=lambda x: x[1], reverse=True)
        return scored_files
    
    def _read_text_preview(self, file_path: Path, max_chars: int = 2000) -> str:
        """Read a preview of a text file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read(max_chars)
            except (UnicodeDecodeError, OSError):
                continue
        
        return ""
    
    def _extract_from_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int = 5000
    ) -> ExtractionResult:
        """
        Extract relevant content from a single file.
        
        Args:
            file_path: Path to the file
            context: Extraction context with objective and keywords
            max_chars: Maximum characters to extract
            
        Returns:
            ExtractionResult with extracted content
        """
        ext = file_path.suffix.lower()
        file_type = self._get_file_type(ext)
        
        try:
            original_size = file_path.stat().st_size
        except OSError:
            original_size = 0
        
        extracted_content = ""
        key_points: List[str] = []
        metadata: Dict[str, Any] = {"file_type": file_type, "extension": ext}
        truncated = False
        
        try:
            if ext in ['.txt', '.md', '.rst', '.log']:
                extracted_content, truncated = self._extract_text_file(
                    file_path, context, max_chars
                )
            elif ext in ['.json']:
                extracted_content, key_points = self._extract_json_file(
                    file_path, context, max_chars
                )
            elif ext in ['.csv']:
                extracted_content, key_points = self._extract_csv_file(
                    file_path, context, max_chars
                )
            elif ext in ['.xml', '.yaml', '.yml']:
                extracted_content, truncated = self._extract_text_file(
                    file_path, context, max_chars
                )
            elif ext == '.pdf' and self.has_pdf:
                extracted_content, key_points = self._extract_pdf_file(
                    file_path, context, max_chars
                )
            elif ext in ['.xlsx', '.xls', '.xlsm'] and self.has_excel:
                extracted_content, key_points = self._extract_excel_file(
                    file_path, context, max_chars
                )
            elif ext in ['.docx'] and self.has_docx:
                extracted_content, key_points = self._extract_docx_file(
                    file_path, context, max_chars
                )
            else:
                extracted_content = f"[File type {ext} - content extraction not supported]"
                metadata["unsupported"] = True
            
        except Exception as e:
            logger.warning(f"Error extracting from {file_path}: {e}")
            extracted_content = f"[Error extracting content: {str(e)}]"
            metadata["error"] = str(e)
        
        # Generate key points if not already done
        if not key_points and extracted_content and not metadata.get("error"):
            key_points = self._extract_key_points(extracted_content, context)
        
        return ExtractionResult(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=file_type,
            relevance_score=0.0,  # Will be set by caller
            extracted_content=extracted_content,
            key_points=key_points,
            metadata=metadata,
            truncated=truncated,
            original_size=original_size,
            extracted_size=len(extracted_content)
        )
    
    def _get_file_type(self, ext: str) -> str:
        """Get file type category from extension."""
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if ext in extensions:
                return file_type
        return "other"
    
    def _extract_text_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int
    ) -> Tuple[str, bool]:
        """Extract content from a text file with relevance filtering."""
        full_text = self._read_text_preview(file_path, self.MAX_TEXT_PREVIEW)
        
        if len(full_text) <= max_chars:
            return full_text, False
        
        # Extract relevant paragraphs
        keywords_lower = [kw.lower() for kw in context.keywords]
        paragraphs = re.split(r'\n\s*\n', full_text)
        
        scored_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            para_lower = para.lower()
            score = sum(1 for kw in keywords_lower if kw in para_lower)
            scored_paragraphs.append((para, score))
        
        # Sort by relevance
        scored_paragraphs.sort(key=lambda x: x[1], reverse=True)
        
        # Build content up to limit
        result_parts = []
        current_size = 0
        
        for para, score in scored_paragraphs:
            if current_size + len(para) + 2 > max_chars:
                break
            result_parts.append(para)
            current_size += len(para) + 2
        
        return "\n\n".join(result_parts), len(full_text) > max_chars
    
    def _extract_json_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int
    ) -> Tuple[str, List[str]]:
        """Extract content from a JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except json.JSONDecodeError:
            # Try to read as text
            content = self._read_text_preview(file_path, max_chars)
            return content, []
        
        key_points = []
        
        # Extract structure summary
        if isinstance(data, dict):
            key_points.append(f"JSON object with keys: {list(data.keys())[:10]}")
            if len(data) > 10:
                key_points.append(f"Total keys: {len(data)}")
        elif isinstance(data, list):
            key_points.append(f"JSON array with {len(data)} items")
            if data and isinstance(data[0], dict):
                key_points.append(f"Item structure: {list(data[0].keys())[:5]}")
        
        # Format JSON with limit
        formatted = json.dumps(data, indent=2, default=str)
        if len(formatted) > max_chars:
            formatted = formatted[:max_chars] + "\n... [truncated]"
        
        return formatted, key_points
    
    def _extract_csv_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int
    ) -> Tuple[str, List[str]]:
        """Extract content from a CSV file."""
        key_points = []
        lines = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                header = f.readline().strip()
                lines.append(header)
                key_points.append(f"Columns: {header}")
                
                row_count = 0
                for line in f:
                    row_count += 1
                    if len('\n'.join(lines)) < max_chars - 200:
                        lines.append(line.strip())
                
                key_points.append(f"Total rows: ~{row_count}")
        except Exception as e:
            return f"[Error reading CSV: {e}]", []
        
        content = '\n'.join(lines)
        if len(content) > max_chars:
            content = content[:max_chars] + "\n... [truncated]"
        
        return content, key_points
    
    def _extract_pdf_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int
    ) -> Tuple[str, List[str]]:
        """Extract content from a PDF file."""
        key_points = []
        text_parts = []
        
        try:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(str(file_path))
                num_pages = len(reader.pages)
            except ImportError:
                import pypdf
                reader = pypdf.PdfReader(str(file_path))
                num_pages = len(reader.pages)
            
            key_points.append(f"PDF with {num_pages} pages")
            
            current_size = 0
            for i, page in enumerate(reader.pages):
                if current_size >= max_chars:
                    break
                    
                text = page.extract_text() or ""
                if text.strip():
                    text_parts.append(f"--- Page {i+1} ---\n{text.strip()}")
                    current_size += len(text) + 20
            
        except Exception as e:
            return f"[Error reading PDF: {e}]", []
        
        content = '\n\n'.join(text_parts)
        if len(content) > max_chars:
            content = content[:max_chars] + "\n... [truncated]"
        
        return content, key_points
    
    def _extract_excel_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int
    ) -> Tuple[str, List[str]]:
        """Extract content from an Excel file."""
        key_points = []
        content_parts = []
        
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            
            key_points.append(f"Excel workbook with sheets: {wb.sheetnames}")
            
            current_size = 0
            for sheet_name in wb.sheetnames:
                if current_size >= max_chars:
                    break
                    
                sheet = wb[sheet_name]
                content_parts.append(f"--- Sheet: {sheet_name} ---")
                
                rows = []
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i >= 50:  # Limit rows
                        rows.append("... [more rows]")
                        break
                    row_str = '\t'.join(str(cell) if cell is not None else '' for cell in row[:20])
                    rows.append(row_str)
                    current_size += len(row_str) + 1
                
                content_parts.append('\n'.join(rows))
            
            wb.close()
            
        except Exception as e:
            return f"[Error reading Excel: {e}]", []
        
        content = '\n\n'.join(content_parts)
        if len(content) > max_chars:
            content = content[:max_chars] + "\n... [truncated]"
        
        return content, key_points
    
    def _extract_docx_file(
        self,
        file_path: Path,
        context: ExtractionContext,
        max_chars: int
    ) -> Tuple[str, List[str]]:
        """Extract content from a Word document."""
        key_points = []
        
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            key_points.append(f"Document with {len(paragraphs)} paragraphs")
            
            # Include tables info
            if doc.tables:
                key_points.append(f"Contains {len(doc.tables)} tables")
            
            content = '\n\n'.join(paragraphs)
            
        except Exception as e:
            return f"[Error reading DOCX: {e}]", []
        
        if len(content) > max_chars:
            content = content[:max_chars] + "\n... [truncated]"
        
        return content, key_points
    
    def _extract_key_points(
        self,
        content: str,
        context: ExtractionContext,
        max_points: int = 5
    ) -> List[str]:
        """Extract key points from content based on context."""
        key_points = []
        keywords_lower = [kw.lower() for kw in context.keywords]
        
        # Split into sentences
        sentences = re.split(r'[.!?]\s+', content)
        
        # Score sentences by keyword matches
        scored = []
        for sent in sentences:
            sent = sent.strip()
            if len(sent) < 20 or len(sent) > 300:
                continue
            sent_lower = sent.lower()
            score = sum(1 for kw in keywords_lower if kw in sent_lower)
            if score > 0:
                scored.append((sent, score))
        
        # Get top sentences as key points
        scored.sort(key=lambda x: x[1], reverse=True)
        for sent, _ in scored[:max_points]:
            key_points.append(sent[:200])
        
        return key_points
    
    def _get_cache_key(self, file_path: Path, context: ExtractionContext) -> str:
        """Generate cache key for an extraction."""
        try:
            mtime = file_path.stat().st_mtime
        except OSError:
            mtime = 0
        
        key_str = f"{file_path}:{mtime}:{context.objective[:100]}:{','.join(context.keywords[:5])}"
        return hashlib.md5(key_str.encode()).hexdigest()
    
    def _build_extraction_summary(
        self,
        extractions: List[ExtractionResult],
        context: ExtractionContext
    ) -> str:
        """Build a summary of all extractions."""
        if not extractions:
            return "No relevant data extracted from input files."
        
        summary_parts = [
            f"Extracted data from {len(extractions)} files relevant to: {context.objective[:100]}"
        ]
        
        # Group by file type
        by_type: Dict[str, List[ExtractionResult]] = {}
        for ext in extractions:
            ft = ext.file_type
            if ft not in by_type:
                by_type[ft] = []
            by_type[ft].append(ext)
        
        for file_type, exts in by_type.items():
            files_list = ", ".join(e.file_name for e in exts[:3])
            if len(exts) > 3:
                files_list += f" (+{len(exts) - 3} more)"
            summary_parts.append(f"- {file_type.upper()}: {files_list}")
        
        # Include top key points
        all_key_points = []
        for ext in extractions:
            all_key_points.extend(ext.key_points[:2])
        
        if all_key_points:
            summary_parts.append("\nKey findings:")
            for point in all_key_points[:5]:
                summary_parts.append(f"  • {point[:150]}")
        
        return "\n".join(summary_parts)
    
    def get_file_preview(
        self,
        file_path: Union[str, Path],
        max_chars: int = 2000
    ) -> Dict[str, Any]:
        """
        Get a quick preview of a single file.
        
        Args:
            file_path: Path to the file
            max_chars: Maximum characters to return
            
        Returns:
            Dictionary with preview content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}
        
        ext = file_path.suffix.lower()
        
        # Create minimal context
        context = ExtractionContext(
            objective="preview",
            keywords=[],
            max_content_per_file=max_chars
        )
        
        extraction = self._extract_from_file(file_path, context, max_chars)
        
        return {
            "success": True,
            "file_name": file_path.name,
            "file_type": extraction.file_type,
            "content": extraction.extracted_content,
            "key_points": extraction.key_points,
            "size_bytes": extraction.original_size,
            "truncated": extraction.truncated
        }
    
    def search_in_files(
        self,
        input_folder: Union[str, Path],
        search_query: str,
        file_types: Optional[List[str]] = None,
        max_results: int = 10
    ) -> Dict[str, Any]:
        """
        Search for specific content across input files.
        
        Args:
            input_folder: Folder to search in
            search_query: Query string to search for
            file_types: Filter to specific file types
            max_results: Maximum number of results
            
        Returns:
            Dictionary with search results
        """
        input_path = Path(input_folder).resolve()
        results = []
        
        # Get files to search
        files = self._scan_folders(input_path, None, file_types)
        
        query_lower = search_query.lower()
        query_words = set(query_lower.split())
        
        for file_path in files:
            ext = file_path.suffix.lower()
            
            # Only search text-based files for now
            if ext not in ['.txt', '.md', '.csv', '.json', '.xml', '.yaml', '.yml', '.log']:
                continue
            
            try:
                content = self._read_text_preview(file_path, 50000)
                content_lower = content.lower()
                
                # Check for matches
                if query_lower in content_lower:
                    # Find matching context
                    idx = content_lower.find(query_lower)
                    start = max(0, idx - 100)
                    end = min(len(content), idx + len(search_query) + 100)
                    snippet = content[start:end]
                    
                    results.append({
                        "file": str(file_path),
                        "file_name": file_path.name,
                        "match_type": "exact",
                        "snippet": f"...{snippet}..."
                    })
                elif query_words & set(content_lower.split()):
                    # Partial word matches
                    matching_words = query_words & set(content_lower.split())
                    results.append({
                        "file": str(file_path),
                        "file_name": file_path.name,
                        "match_type": "partial",
                        "matching_words": list(matching_words)
                    })
                
                if len(results) >= max_results:
                    break
                    
            except Exception as e:
                logger.debug(f"Error searching {file_path}: {e}")
        
        return {
            "success": True,
            "query": search_query,
            "results_count": len(results),
            "results": results
        }
    
    def summarize_input_folder(
        self,
        input_folder: Union[str, Path],
        objective: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Get a high-level summary of the input folder contents.
        
        Args:
            input_folder: Path to input folder
            objective: Optional objective for relevance hints
            
        Returns:
            Dictionary with folder summary
        """
        input_path = Path(input_folder).resolve()
        
        if not input_path.exists():
            return {"success": False, "error": f"Folder not found: {input_folder}"}
        
        # Scan files
        all_files = self._scan_folders(input_path, None, None)
        
        # Categorize
        by_type: Dict[str, List[Dict]] = {}
        total_size = 0
        
        for file_path in all_files:
            ext = file_path.suffix.lower()
            file_type = self._get_file_type(ext)
            
            if file_type not in by_type:
                by_type[file_type] = []
            
            try:
                size = file_path.stat().st_size
            except OSError:
                size = 0
            
            by_type[file_type].append({
                "name": file_path.name,
                "size": size,
                "path": str(file_path)
            })
            total_size += size
        
        # Build summary
        summary = {
            "success": True,
            "folder": str(input_path),
            "total_files": len(all_files),
            "total_size_bytes": total_size,
            "total_size_readable": self._format_size(total_size),
            "by_type": {}
        }
        
        for file_type, files in by_type.items():
            type_size = sum(f["size"] for f in files)
            summary["by_type"][file_type] = {
                "count": len(files),
                "size_bytes": type_size,
                "size_readable": self._format_size(type_size),
                "files": [f["name"] for f in files[:10]]  # First 10 names
            }
            if len(files) > 10:
                summary["by_type"][file_type]["more_files"] = len(files) - 10
        
        # Add relevance hints if objective provided
        if objective:
            keywords = self._extract_keywords(objective)
            relevant_files = []
            
            for file_path in all_files[:50]:  # Check first 50
                name_lower = file_path.stem.lower()
                if any(kw in name_lower for kw in keywords):
                    relevant_files.append(file_path.name)
            
            summary["relevance_hints"] = {
                "keywords_extracted": keywords[:10],
                "potentially_relevant_files": relevant_files[:10]
            }
        
        return summary
    
    def _format_size(self, size_bytes: int) -> str:
        """Format bytes to human-readable size."""
        size = float(size_bytes)
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} TB"
    
    def clear_cache(self):
        """Clear the extraction cache."""
        self._extraction_cache.clear()
        logger.info("Extraction cache cleared")
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Get statistics about the extraction cache."""
        return {
            "session_id": self._session_id,
            "cached_extractions": len(self._extraction_cache),
            "total_cached_content": sum(
                len(e.extracted_content) for e in self._extraction_cache.values()
            )
        }


# Required for os.walk
import os
