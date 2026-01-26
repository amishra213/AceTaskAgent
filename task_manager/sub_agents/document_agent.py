"""
Document Sub-Agent for handling text and Word document operations.

Capabilities:
- Create and write TXT files with formatted content
- Create and write DOCX files with formatted content
- Read TXT and DOCX files
- Support for headings, paragraphs, lists, and basic formatting
- Append content to existing documents

Migration Status: Week 7 Day 1 - Dual Format Support
- Supports both legacy dict and standardized AgentExecutionRequest/Response
- Maintains 100% backward compatibility
- Publishes SystemEvent on completion for event-driven workflows
"""

from typing import Optional, List, Dict, Any, Union
from pathlib import Path
from datetime import datetime
import json
import time

from task_manager.utils.logger import get_logger

# Import standardized schemas and utilities
from task_manager.models import (
    AgentExecutionRequest,
    AgentExecutionResponse,
    create_system_event,
    create_error_response
)
from task_manager.utils import (
    auto_convert_response,
    validate_agent_execution_response,
    exception_to_error_response,
    InvalidParameterError,
    wrap_exception
)
from task_manager.core.event_bus import get_event_bus

logger = get_logger(__name__)


class DocumentAgent:
    """
    Sub-agent for text and Word document operations.
    
    This agent handles all document-related tasks:
    - Creating and writing TXT files
    - Creating and writing DOCX files
    - Reading document content
    - Formatting text with headings, paragraphs, and lists
    - Appending to existing documents
    """
    
    def __init__(self):
        """Initialize Document Agent with dual-format support."""
        self.agent_name = "document_agent"
        self.supported_operations = [
            "create_txt",
            "create_docx",
            "read_txt",
            "read_docx",
            "append_txt",
            "append_docx"
        ]
        
        # Initialize event bus for event-driven workflows
        self.event_bus = get_event_bus()
        
        logger.info("Document Agent initialized with dual-format support")
        self._check_dependencies()
    
    
    def _check_dependencies(self):
        """Check if required document libraries are installed."""
        try:
            import docx
            self.docx_lib = "python-docx"
            logger.debug("Using python-docx for DOCX operations")
        except ImportError:
            logger.warning(
                "python-docx not found. Install with: pip install python-docx"
            )
            self.docx_lib = None
    
    
    def _determine_output_path(
        self,
        parameters: Dict[str, Any],
        default_extension: str,
        operation: str
    ) -> str:
        """
        Determine output path from various parameter combinations.
        
        Supports multiple ways to specify output location:
        1. Direct: output_path or file_path
        2. Combined: folder_path + file_name
        3. Template-based: template_name (uses current directory)
        4. Title-based: title (uses current directory)
        
        Args:
            parameters: Task parameters dictionary
            default_extension: File extension to use (e.g., 'txt', 'docx')
            operation: Operation name for logging
        
        Returns:
            Resolved output path as string
        """
        output_path_raw = parameters.get('output_path')
        file_path_raw = parameters.get('file_path')
        folder_path = parameters.get('folder_path')
        file_name = parameters.get('file_name')
        template_name = parameters.get('template_name')
        title = parameters.get('title')
        
        logger.info(f"[DOCUMENT AGENT] Raw output_path: '{output_path_raw}'")
        logger.info(f"[DOCUMENT AGENT] Raw file_path: '{file_path_raw}'")
        logger.info(f"[DOCUMENT AGENT] Folder path: '{folder_path}'")
        logger.info(f"[DOCUMENT AGENT] File name: '{file_name}'")
        logger.info(f"[DOCUMENT AGENT] Template name: '{template_name}'")
        logger.info(f"[DOCUMENT AGENT] Title: '{title}'")
        
        # Build output path from various sources
        output_path = output_path_raw or file_path_raw
        
        if not output_path and folder_path and file_name:
            # Construct the path from folder and filename
            # Add extension if not present
            if not file_name.endswith(f'.{default_extension}'):
                file_name = f"{file_name}.{default_extension}"
            output_path = str(Path(folder_path) / file_name)
            logger.info(f"[DOCUMENT AGENT] Constructed output_path from folder_path + file_name: '{output_path}'")
        
        elif not output_path and folder_path and template_name:
            # Use template name with folder
            safe_name = template_name.replace(' ', '_').replace('Template', '').strip('_')
            output_path = str(Path(folder_path) / f"{safe_name}.{default_extension}")
            logger.info(f"[DOCUMENT AGENT] Constructed output_path from folder_path + template_name: '{output_path}'")
        
        elif not output_path and folder_path and title:
            # Use title with folder
            safe_name = title.replace(' ', '_')
            output_path = str(Path(folder_path) / f"{safe_name}.{default_extension}")
            logger.info(f"[DOCUMENT AGENT] Constructed output_path from folder_path + title: '{output_path}'")
        
        elif not output_path and template_name:
            # Use template name in current directory
            safe_name = template_name.replace(' ', '_').replace('Template', '').strip('_')
            output_path = str(Path('output_folder') / f"{safe_name}.{default_extension}")
            logger.warning(f"[DOCUMENT AGENT] No folder specified, using output_folder: '{output_path}'")
        
        elif not output_path and title:
            # Use title in current directory
            safe_name = title.replace(' ', '_')
            output_path = str(Path('output_folder') / f"{safe_name}.{default_extension}")
            logger.warning(f"[DOCUMENT AGENT] No folder specified, using output_folder: '{output_path}'")
        
        elif not output_path:
            # Last resort: use operation name with timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = str(Path('output_folder') / f"{operation}_{timestamp}.{default_extension}")
            logger.warning(f"[DOCUMENT AGENT] No path info provided, using timestamped name: '{output_path}'")
        
        logger.info(f"[DOCUMENT AGENT] Final output_path for {operation}: '{output_path}'")
        return output_path
    
    
    def create_txt(
        self,
        output_path: str,
        content: Union[str, List[str]],
        encoding: str = "utf-8"
    ) -> Dict[str, Any]:
        """
        Create a TXT file with the specified content.
        
        Args:
            output_path: Path where the TXT file will be created
            content: String or list of strings (lines) to write
            encoding: Text encoding (default: utf-8)
        
        Returns:
            Dictionary with operation result
        """
        try:
            output_path_obj = Path(output_path)
            
            # Ensure parent directory exists
            output_path_obj.parent.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Creating TXT file: {output_path}")
            
            # Convert content to string
            if isinstance(content, list):
                text_content = "\n".join(str(line) for line in content)
            else:
                text_content = str(content)
            
            # Write file
            with open(output_path_obj, 'w', encoding=encoding) as f:
                f.write(text_content)
            
            file_size = output_path_obj.stat().st_size
            
            logger.info(f"TXT file created successfully: {output_path} ({file_size} bytes)")
            
            return {
                "success": True,
                "file": str(output_path),
                "output_path": str(output_path),
                "size_bytes": file_size,
                "lines": len(text_content.split('\n')),
                "encoding": encoding,
                "created_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error creating TXT file: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file": str(output_path)
            }
    
    
    def create_docx(
        self,
        output_path: str,
        content: Union[str, List[Dict[str, Any]]],
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create a DOCX file with formatted content.
        
        Args:
            output_path: Path where the DOCX file will be created
            content: String or list of content blocks with formatting
                     Each block can be: {'type': 'heading'|'paragraph'|'list', 'text': '...', 'level': 1}
            title: Optional document title
        
        Returns:
            Dictionary with operation result
        """
        try:
            if not self.docx_lib:
                return {
                    "success": False,
                    "error": "python-docx library not installed. Install with: pip install python-docx",
                    "file": output_path
                }
            
            from docx import Document
            from docx.shared import Pt, RGBColor
            
            output_path_obj = Path(output_path)
            
            # Ensure parent directory exists
            output_path_obj.parent.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Creating DOCX file: {output_path}")
            
            # Create document
            doc = Document()
            
            # Add title if provided
            if title:
                heading = doc.add_heading(title, level=0)
            
            # Process content
            if isinstance(content, str):
                # Simple string content
                doc.add_paragraph(content)
                blocks_added = 1
            elif isinstance(content, list):
                blocks_added = 0
                for block in content:
                    if isinstance(block, str):
                        doc.add_paragraph(block)
                        blocks_added += 1
                    elif isinstance(block, dict):
                        block_type = block.get('type', 'paragraph')
                        text = block.get('text', '')
                        
                        if block_type == 'heading':
                            level = block.get('level', 1)
                            doc.add_heading(text, level=level)
                            blocks_added += 1
                        elif block_type == 'paragraph':
                            p = doc.add_paragraph(text)
                            # Apply formatting if specified
                            if block.get('bold'):
                                p.runs[0].bold = True
                            if block.get('italic'):
                                p.runs[0].italic = True
                            blocks_added += 1
                        elif block_type == 'list':
                            items = block.get('items', [text])
                            for item in items:
                                doc.add_paragraph(str(item), style='List Bullet')
                                blocks_added += 1
            else:
                doc.add_paragraph(str(content))
                blocks_added = 1
            
            # Save document
            doc.save(str(output_path_obj))
            
            file_size = output_path_obj.stat().st_size
            
            logger.info(f"DOCX file created successfully: {output_path} ({file_size} bytes)")
            
            return {
                "success": True,
                "file": str(output_path),
                "output_path": str(output_path),
                "size_bytes": file_size,
                "blocks_added": blocks_added,
                "title": title,
                "created_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error creating DOCX file: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file": str(output_path)
            }
    
    
    def read_txt(
        self,
        file_path: str,
        encoding: str = "utf-8"
    ) -> Dict[str, Any]:
        """
        Read content from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            encoding: Text encoding (default: utf-8)
        
        Returns:
            Dictionary with file content
        """
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return {
                    "success": False,
                    "error": f"File not found: {file_path}",
                    "file": str(file_path_obj)
                }
            
            logger.info(f"Reading TXT file: {file_path}")
            
            with open(file_path_obj, 'r', encoding=encoding) as f:
                content = f.read()
            
            lines = content.split('\n')
            
            return {
                "success": True,
                "file": str(file_path),
                "content": content,
                "lines": lines,
                "line_count": len(lines),
                "char_count": len(content),
                "encoding": encoding,
                "read_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error reading TXT file: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file": str(file_path)
            }
    
    
    def read_docx(
        self,
        file_path: str
    ) -> Dict[str, Any]:
        """
        Read content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
        
        Returns:
            Dictionary with file content
        """
        try:
            if not self.docx_lib:
                return {
                    "success": False,
                    "error": "python-docx library not installed",
                    "file": file_path
                }
            
            from docx import Document
            
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return {
                    "success": False,
                    "error": f"File not found: {file_path}",
                    "file": str(file_path_obj)
                }
            
            logger.info(f"Reading DOCX file: {file_path}")
            
            doc = Document(str(file_path_obj))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })
            
            # Extract full text
            full_text = "\n".join(para.text for para in doc.paragraphs)
            
            return {
                "success": True,
                "file": str(file_path),
                "paragraphs": paragraphs,
                "full_text": full_text,
                "paragraph_count": len(paragraphs),
                "char_count": len(full_text),
                "read_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file": str(file_path)
            }
    
    
    def append_txt(
        self,
        file_path: str,
        content: Union[str, List[str]],
        encoding: str = "utf-8"
    ) -> Dict[str, Any]:
        """
        Append content to an existing TXT file.
        
        Args:
            file_path: Path to the TXT file
            content: String or list of strings to append
            encoding: Text encoding (default: utf-8)
        
        Returns:
            Dictionary with operation result
        """
        try:
            file_path_obj = Path(file_path)
            
            # Ensure parent directory exists
            file_path_obj.parent.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Appending to TXT file: {file_path}")
            
            # Convert content to string
            if isinstance(content, list):
                text_content = "\n".join(str(line) for line in content)
            else:
                text_content = str(content)
            
            # Append to file
            with open(file_path_obj, 'a', encoding=encoding) as f:
                f.write("\n" + text_content)
            
            file_size = file_path_obj.stat().st_size
            
            logger.info(f"Content appended successfully to: {file_path}")
            
            return {
                "success": True,
                "file": str(file_path),
                "size_bytes": file_size,
                "appended_lines": len(text_content.split('\n')),
                "encoding": encoding,
                "appended_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error appending to TXT file: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file": str(file_path)
            }
    
    
    def append_docx(
        self,
        file_path: str,
        content: Union[str, List[Dict[str, Any]]]
    ) -> Dict[str, Any]:
        """
        Append content to an existing DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            content: String or list of content blocks to append
        
        Returns:
            Dictionary with operation result
        """
        try:
            if not self.docx_lib:
                return {
                    "success": False,
                    "error": "python-docx library not installed",
                    "file": file_path
                }
            
            from docx import Document
            
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return {
                    "success": False,
                    "error": f"File not found: {file_path}",
                    "file": str(file_path_obj)
                }
            
            logger.info(f"Appending to DOCX file: {file_path}")
            
            # Open existing document
            doc = Document(str(file_path_obj))
            
            # Process content
            if isinstance(content, str):
                doc.add_paragraph(content)
                blocks_added = 1
            elif isinstance(content, list):
                blocks_added = 0
                for block in content:
                    if isinstance(block, str):
                        doc.add_paragraph(block)
                        blocks_added += 1
                    elif isinstance(block, dict):
                        block_type = block.get('type', 'paragraph')
                        text = block.get('text', '')
                        
                        if block_type == 'heading':
                            level = block.get('level', 1)
                            doc.add_heading(text, level=level)
                            blocks_added += 1
                        elif block_type == 'paragraph':
                            doc.add_paragraph(text)
                            blocks_added += 1
                        elif block_type == 'list':
                            items = block.get('items', [text])
                            for item in items:
                                doc.add_paragraph(str(item), style='List Bullet')
                                blocks_added += 1
            else:
                doc.add_paragraph(str(content))
                blocks_added = 1
            
            # Save document
            doc.save(str(file_path_obj))
            
            file_size = file_path_obj.stat().st_size
            
            logger.info(f"Content appended successfully to: {file_path}")
            
            return {
                "success": True,
                "file": str(file_path),
                "size_bytes": file_size,
                "blocks_added": blocks_added,
                "appended_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error appending to DOCX file: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file": str(file_path)
            }
    
    
    def execute_task(
        self,
        *args: Any,
        **kwargs: Any
    ) -> Union[Dict[str, Any], AgentExecutionResponse]:
        """
        Execute a document operation with dual-format support.
        
        Supports three calling conventions:
        1. Legacy positional: execute_task(operation, parameters)
        2. Legacy dict: execute_task({'operation': ..., 'parameters': ...})
        3. Standardized: execute_task(AgentExecutionRequest)
        
        Args:
            *args: Variable positional arguments
            **kwargs: Variable keyword arguments
        
        Returns:
            Legacy dict OR AgentExecutionResponse based on input format
        """
        start_time = time.time()
        return_legacy = True
        operation = None
        parameters = None
        task_dict = None
        
        # Detect calling convention
        if len(args) == 2:
            operation, parameters = args
            task_dict = {"operation": operation, "parameters": parameters}
            return_legacy = True
            logger.debug("Legacy positional call")
        
        elif len(args) == 1 and isinstance(args[0], dict):
            task_dict = args[0]
            if "task_id" in task_dict and "task_description" in task_dict:
                return_legacy = False
                logger.debug(f"Standardized request call: task_id={task_dict.get('task_id')}")
            else:
                return_legacy = True
                logger.debug("Legacy dict call")
            operation = task_dict.get("operation")
            parameters = task_dict.get("parameters", {})
        
        elif "operation" in kwargs:
            operation = kwargs.get("operation")
            parameters = kwargs.get("parameters", {})
            task_dict = {"operation": operation, "parameters": parameters}
            return_legacy = True
            logger.debug("Legacy keyword call")
        
        else:
            raise InvalidParameterError(
                parameter_name="task",
                message="Invalid call to execute_task. Use one of:\n"
                "  - execute_task(operation, parameters)\n"
                "  - execute_task({'operation': ..., 'parameters': ...})\n"
                "  - execute_task(AgentExecutionRequest)"
            )
        
        try:
            task_id = task_dict.get("task_id", f"document_{int(time.time())}")
            
            if parameters is None:
                parameters = {}
            
            if operation is None:
                operation = "unknown"
            
            logger.info("=" * 80)
            logger.info(f"[DOCUMENT AGENT] Starting execution")
            logger.info("=" * 80)
            logger.info(f"[DOCUMENT AGENT] Operation: {operation}")
            logger.info(f"[DOCUMENT AGENT] Task ID: {task_id}")
            logger.info(f"[DOCUMENT AGENT] Return format: {'legacy' if return_legacy else 'standardized'}")
            logger.info(f"[DOCUMENT AGENT] Parameters: {json.dumps(parameters, indent=2, default=str)}")
            
            # Execute the operation
            if operation == "create_txt":
                output_path = self._determine_output_path(parameters, 'txt', operation)
                result = self.create_txt(
                    output_path=output_path,
                    content=parameters.get('content', ''),
                    encoding=parameters.get('encoding', 'utf-8')
                )
            
            elif operation == "create_docx":
                output_path = self._determine_output_path(parameters, 'docx', operation)
                result = self.create_docx(
                    output_path=output_path,
                    content=parameters.get('content', ''),
                    title=parameters.get('title')
                )
            
            elif operation == "read_txt":
                result = self.read_txt(
                    file_path=parameters.get('file_path', ''),
                    encoding=parameters.get('encoding', 'utf-8')
                )
            
            elif operation == "read_docx":
                result = self.read_docx(
                    file_path=parameters.get('file_path', '')
                )
            
            elif operation == "append_txt":
                result = self.append_txt(
                    file_path=parameters.get('file_path', ''),
                    content=parameters.get('content', ''),
                    encoding=parameters.get('encoding', 'utf-8')
                )
            
            elif operation == "append_docx":
                result = self.append_docx(
                    file_path=parameters.get('file_path', ''),
                    content=parameters.get('content', '')
                )
            
            else:
                result = {
                    "success": False,
                    "error": f"Unsupported operation: {operation}",
                    "supported_operations": self.supported_operations
                }
            
            logger.info(f"[DOCUMENT AGENT] Operation result: success={result.get('success')}")
            
            # Convert to standardized format if needed
            if return_legacy:
                response = result
            else:
                response = self._convert_to_standard_response(result, operation, task_id, start_time)
                self._publish_completion_event(task_id, operation, response)
            
            return response
        
        except Exception as e:
            logger.error(f"Error in execute_task: {str(e)}", exc_info=True)
            
            error = exception_to_error_response(
                e,
                source=self.agent_name,
                task_id=task_dict.get("task_id", "unknown") if task_dict else "unknown"
            )
            
            error_response: AgentExecutionResponse = {
                "status": "failure",
                "success": False,
                "result": {},
                "artifacts": [],
                "execution_time_ms": int((time.time() - start_time) * 1000),
                "timestamp": datetime.now().isoformat(),
                "agent_name": self.agent_name,
                "operation": operation or "unknown",
                "blackboard_entries": [],
                "warnings": []
            }
            # Add error field separately to handle TypedDict
            error_response["error"] = error  # type: ignore
            
            if return_legacy:
                return self._convert_to_legacy_response(error_response)
            else:
                return error_response
    
    
    def _convert_to_standard_response(
        self,
        legacy_result: Dict[str, Any],
        operation: str,
        task_id: str,
        start_time: float
    ) -> AgentExecutionResponse:
        """Convert legacy result dict to standardized AgentExecutionResponse."""
        success = legacy_result.get("success", False)
        
        # Extract artifacts from result
        artifacts = []
        output_file = legacy_result.get("file") or legacy_result.get("output_path")
        if output_file and success:
            file_path = Path(output_file)
            if file_path.exists():
                file_ext = file_path.suffix.lower()
                artifacts.append({
                    "type": file_ext[1:] if file_ext else "txt",
                    "path": str(output_file),
                    "size_bytes": file_path.stat().st_size,
                    "description": f"Document output from {operation} operation"
                })
        
        # Create blackboard entries for sharing data
        blackboard_entries = []
        if success and "content" in legacy_result:
            blackboard_entries.append({
                "key": f"document_content_{task_id}",
                "value": legacy_result.get("content", ""),
                "scope": "workflow",
                "ttl_seconds": 3600
            })
        
        # Build standardized response
        response: AgentExecutionResponse = {
            "status": "success" if success else "failure",
            "success": success,
            "result": {
                k: v for k, v in legacy_result.items()
                if k not in ["success", "file", "output_path"]
            },
            "artifacts": artifacts,
            "execution_time_ms": int((time.time() - start_time) * 1000),
            "timestamp": datetime.now().isoformat(),
            "agent_name": self.agent_name,
            "operation": operation,
            "blackboard_entries": blackboard_entries,
            "warnings": []
        }
        
        # Add error field if present (handle TypedDict)
        if not success and "error" in legacy_result:
            response["error"] = create_error_response(  # type: ignore
                error_code="DOC_001",
                error_type="execution_error",
                message=legacy_result.get("error", "Unknown error"),
                source=self.agent_name
            )
        
        return response
    
    
    def _convert_to_legacy_response(self, standard_response: AgentExecutionResponse) -> Dict[str, Any]:
        """Convert standardized response back to legacy format for backward compatibility."""
        legacy = {
            "success": standard_response["success"],
            "result": standard_response["result"]
        }
        
        if standard_response["artifacts"]:
            legacy["file"] = standard_response["artifacts"][0]["path"]
            legacy["output_path"] = standard_response["artifacts"][0]["path"]
        
        # Add error if present (use .get() for NotRequired field)
        error = standard_response.get("error")  # type: ignore
        if error:
            legacy["error"] = error["message"]  # type: ignore
        
        if isinstance(standard_response["result"], dict):
            for key, value in standard_response["result"].items():
                if key not in legacy:
                    legacy[key] = value
        
        return legacy
    
    
    def _publish_completion_event(
        self,
        task_id: str,
        operation: str,
        response: AgentExecutionResponse
    ):
        """Publish task completion event for event-driven workflows."""
        try:
            event = create_system_event(
                event_type="document_created" if "create" in operation else "document_processed",
                event_category="task_lifecycle",
                source_agent=self.agent_name,
                payload={
                    "task_id": task_id,
                    "operation": operation,
                    "success": response["success"],
                    "artifacts": response["artifacts"],
                    "blackboard_keys": [entry["key"] for entry in response["blackboard_entries"]]
                }
            )
            self.event_bus.publish(event)
            logger.debug(f"Published completion event for task {task_id}")
        except Exception as e:
            logger.warning(f"Failed to publish completion event: {e}")
